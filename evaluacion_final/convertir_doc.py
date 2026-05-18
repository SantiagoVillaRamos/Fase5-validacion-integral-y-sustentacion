#!/usr/bin/env python3
"""
Convertidor de Markdown a Word (.docx) con Estilo APA 7.
Toma el reporte académico '01_reporte_academico_fase5.md'
y genera 'Fase5Santiago_Villa.docx' con tipografía y espaciado APA.
"""

import os
import re
import sys

# Asegurar que python-docx esté instalado en el entorno
try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("El módulo 'python-docx' no está instalado. Instalándolo en el entorno virtual...")
    os.system(f"{sys.executable} -m pip install python-docx")
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

# Rutas de archivos
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
MD_PATH = os.path.join(BASE_DIR, "../docs/01_reporte_academico_fase5.md")
DOCX_PATH = os.path.join(BASE_DIR, "../docs/Fase5Santiago_Villa.docx")


def set_cell_border(cell, **kwargs):
    """Establece los bordes de una celda de tabla en python-docx."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, val in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(val))


def parse_inline_formatting(paragraph, text, center=False):
    """Parsea formatos inline básicos como **negrita** e *itálica* a runs de python-docx."""
    # Encontrar patrones de negrita y cursiva
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|[^\*`]+)')
    tokens = pattern.findall(text)
    
    for token in tokens:
        run = paragraph.add_run()
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x11, 0x18, 0x27) # Dark Charcoal
        
        if token.startswith('**') and token.endswith('**'):
            run.text = token[2:-2]
            run.bold = True
        elif token.startswith('*') and token.endswith('*'):
            run.text = token[1:-1]
            run.italic = True
        elif token.startswith('`') and token.endswith('`'):
            run.text = token[1:-1]
            run.font.name = 'Courier New'
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED) # Accent purple for inline code
        else:
            # Reemplazos HTML comunes
            clean_text = token.replace('<br>', '\n').replace('<br/>', '\n')
            run.text = clean_text


def convertir_md_a_docx():
    print(f"Leyendo documento Markdown en: {MD_PATH}")
    if not os.path.exists(MD_PATH):
        print(f"Error: No se encontró el archivo en {MD_PATH}")
        sys.exit(1)

    with open(MD_PATH, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # Inicializar documento Word
    doc = docx.Document()
    
    # ── Configurar Margenes APA 7 (1 pulgada = 2.54 cm en todos los lados) ──
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Variables de control
    in_yaml = False
    yaml_lines_passed = 0
    in_center_div = False
    
    print("Procesando y aplicando estilos APA 7...")

    for line in lines:
        line_strip = line.strip()
        
        # Omitir el bloque de metadatos YAML
        if line_strip == '---':
            yaml_lines_passed += 1
            in_yaml = (yaml_lines_passed % 2 != 0)
            continue
        if in_yaml:
            continue

        # Detectar contenedores de alineación HTML del Cover Page
        if '<div align="center">' in line_strip:
            in_center_div = True
            continue
        if '</div>' in line_strip:
            in_center_div = False
            continue

        # Líneas horizontales de separación en MD -> Saltos de página en Word para APA
        if line_strip in ['---', '***', '___']:
            doc.add_page_break()
            continue

        # ── Encabezados (H1, H2, H3) ──
        if line_strip.startswith('# '):
            text = line_strip[2:]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(16)
            run.bold = True
            run.font.color.rgb = RGBColor(0x0A, 0x0F, 0x1E)
            continue

        elif line_strip.startswith('## '):
            text = line_strip[3:]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x6F, 0xEB)
            continue

        elif line_strip.startswith('### '):
            text = line_strip[4:]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
            run.italic = True
            run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
            continue

        # ── Listas con Viñetas ──
        elif line_strip.startswith('* ') or line_strip.startswith('- '):
            text = line_strip[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            parse_inline_formatting(p, text)
            continue

        # ── Párrafos regulares o Cover Page ──
        elif line_strip:
            p = doc.add_paragraph()
            
            # Espaciado APA 7 (Interlineado 1.5)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(8)
            
            # Alineación
            if in_center_div:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Sangría extra si es portada para espaciar elementos
                if line_strip.startswith('\\') or line_strip == '':
                    p.paragraph_format.space_before = Pt(18)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            parse_inline_formatting(p, line_strip)
        else:
            # Línea vacía: añadir espaciado al párrafo anterior
            pass

    # Guardar documento
    doc.save(DOCX_PATH)
    print(f"\n¡Éxito! El documento de Word ha sido creado en: {DOCX_PATH}")


if __name__ == "__main__":
    convertir_md_a_docx()
